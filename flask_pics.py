from flask import Flask, request, redirect, flash
import pandas as pd
import requests
import shutil
import docx
from docx.shared import Inches

app = Flask(__name__)


@app.route('/', methods=['GET', 'POST'])
def read_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file')
            return redirect(request.url)
        file = request.files['file']
        df = pd.read_csv(file, header=None)
        pic_list = df.values.tolist()
        doc = docx.Document()
        doc.add_heading('TEST!!!', 0)
        for pic in pic_list:
            print(pic[0])
            picture_obj = requests.get(fr'{pic[0]}', stream=True)
            picture_obj.raw.decode_content = True
            with open("pic_saveover.jpg", "wb") as f:
                shutil.copyfileobj(picture_obj.raw, f)
            doc.add_picture('pic_saveover.jpg', width=Inches(2.8))
        doc.save('test.docx')
        return 'File Saved!'

    return '''
    <!doctype html>
    <title>Upload filepath</title>
    <h1>Upload filepath</h1>
    <form method=post enctype=multipart/form-data>
        <input type=file name=file>
        <input type=submit value=Upload>
    </form>
    '''
while __name__ == '__main__':
    app.run(debug=True)